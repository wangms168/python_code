# from datetime import date
from time import sleep
from openpyxl import load_workbook
from dateutil.relativedelta import relativedelta
import os

import pandas as pd

pd.set_option('display.unicode.ambiguous_as_wide', True)
pd.set_option('display.unicode.east_asian_width', True)
pd.set_option('display.width', 180)

# ---------------------------------------对比分析--------------------------------------------
xlfile = "对比分析.xlsx"
wb = load_workbook(xlfile)
# ws = wb.active
ws = wb['科目余额表-母公司']
date_本年当月 = ws.cell(row=3, column=6).value  # 获取单元格的内容
date_本年上月 = date_本年当月 + relativedelta(months=-1)
date_上年当月 = date_本年当月 + relativedelta(years=-1)
date_上年上月 = date_上年当月 + relativedelta(months=-1)

date_本年当月月份 = date_本年当月.strftime('%m')
date_本年当月 = date_本年当月.strftime('%Y%m')
date_本年上月 = date_本年上月.strftime('%Y%m')
date_上年当月 = date_上年当月.strftime('%Y%m')
date_上年上月 = date_上年上月.strftime('%Y%m')

FilesName_本年上月 = '科目余额表-母公司-' + date_本年上月 + '.xls'
FilesName_本年当月 = '科目余额表-母公司-' + date_本年当月 + '.xls'
FilesName_上年当月 = '科目余额表-母公司-' + date_上年当月 + '.xls'

# print("本年当月_文件名:", FilesName_本年当月)
# print("本年上月_文件名:", FilesName_本年上月)
# print("上年当月_文件名:", FilesName_上年当月)

df_分析xl = pd.read_excel(xlfile, skiprows=4 - 1)
list_科目名称 = df_分析xl['科目名称']

# print(list_科目名称.index)
# print(list_科目名称.values)
# print(list_科目名称.keys)
# print(list_科目名称.items)
# for key in list_科目名称.keys():
#     print(key)
# for key_value in list_科目名称.items():
#     print(key_value)
# print(type(list_科目名称), list_科目名称)
# print(list_科目名称.isnull())


# -----------------------------------------科目余额表--------------------------------------------
for key_value in list_科目名称.notna().items():
    if key_value[1]:
        row = key_value[0] + 5
        科目名称 = list_科目名称.values[key_value[0]]
        # print(row, 科目名称)

        if not os.path.exists(FilesName_本年上月):  # 符合条件的excel不存在，程序退出
            print(FilesName_本年上月 + "不存在")
        else:
            df_本年当月 = pd.read_excel(FilesName_本年上月, index_col=1 - 1, skiprows=9 - 1)
            df_本年当月 = df_本年当月.loc[:"总计", :]  # df切片，截取0到"合计"连续的行、所有列
            df_本年当月.reset_index(["科目编码"], inplace=True)  # 取消 原in_df 以"xm"列作为行索引的设置
            df_本年当月.set_index(["科目名称"], inplace=True)  # 设置 dd_df 以 "代垫" 列 作为含索引

            本年当月_本期贷方 = df_本年当月['本币.2'][科目名称]  # '本币.2'本期贷方
            本年当月_本年累计 = df_本年当月['本币.4'][科目名称]  # '本币.4'贷方累计

            ws.cell(row=row, column=4).value = 本年当月_本期贷方
            ws.cell(row=row, column=5).value = 本年当月_本年累计
            wb.save(filename=xlfile)

        if not os.path.exists(FilesName_本年当月):  # 符合条件的excel不存在，程序退出
            print(FilesName_本年当月 + "不存在")
        else:
            df_本年当月 = pd.read_excel(FilesName_本年当月, index_col=1 - 1, skiprows=9 - 1)
            df_本年当月 = df_本年当月.loc[:"总计", :]  # df切片，截取0到"合计"连续的行、所有列
            df_本年当月.reset_index(["科目编码"], inplace=True)  # 取消 原in_df 以"xm"列作为行索引的设置
            df_本年当月.set_index(["科目名称"], inplace=True)  # 设置 dd_df 以 "代垫" 列 作为含索引

            本年当月_本期贷方 = df_本年当月['本币.2'][科目名称]  # '本币.2'本期贷方
            本年当月_本年累计 = df_本年当月['本币.4'][科目名称]  # '本币.4'贷方累计

            ws.cell(row=row, column=6).value = 本年当月_本期贷方
            ws.cell(row=row, column=7).value = 本年当月_本年累计
            wb.save(filename=xlfile)

        if not os.path.exists(FilesName_上年当月):  # 符合条件的excel不存在，程序退出
            print(FilesName_上年当月 + "不存在")
        else:
            df_本年当月 = pd.read_excel(FilesName_上年当月, index_col=1 - 1, skiprows=9 - 1)
            df_本年当月 = df_本年当月.loc[:"总计", :]  # df切片，截取0到"合计"连续的行、所有列
            df_本年当月.reset_index(["科目编码"], inplace=True)  # 取消 原in_df 以"xm"列作为行索引的设置
            df_本年当月.set_index(["科目名称"], inplace=True)  # 设置 dd_df 以 "代垫" 列 作为含索引

            本年当月_本期贷方 = df_本年当月['本币.2'][科目名称]  # '本币.2'本期贷方
            本年当月_本年累计 = df_本年当月['本币.4'][科目名称]  # '本币.4'贷方累计

            ws.cell(row=row, column=8).value = 本年当月_本期贷方
            ws.cell(row=row, column=9).value = 本年当月_本年累计
            wb.save(filename=xlfile)
wb.close()

from win32com.client import Dispatch

xlApp = Dispatch("Excel.Application")
xlApp.Visible = False
xlfile = os.path.abspath(xlfile)  # win32不认识相对路径，故需转换为绝对路径。
xlBook = xlApp.Workbooks.Open(xlfile)
xlBook.Save()
xlBook.Close()

# ---------------------------------------辅助余额表-总部-部门--------------------------------------------
wb = load_workbook(xlfile, data_only=True)
ws = wb.active

Unit = ws.cell(row=2, column=2).value
f21 = ws.cell(row=21, column=6).value
F21 = "{:,}".format(round(f21, 2)) if f21 else ''
j21 = ws.cell(row=21, column=10).value
J21 = "{:,}".format(round(j21, 2)) if j21 else ''
if j21 is None or f21 == 0:
    J21_ad = '持平'
    J21_ll = '平'
elif j21 > 0.0:
    J21_ad = '增加'
    J21_ll = '升'
elif j21 < 0.0:
    J21_ad = '减少'
    J21_ll = '降'
k21 = ws.cell(row=21, column=11).value
K21 = "{:,}".format(round(k21, 2)) if k21 else ''

f22 = ws.cell(row=22, column=6).value
F22 = "{:,}".format(round(f22, 2)) if f22 else ''
j22 = ws.cell(row=22, column=10).value
J22 = "{:,}".format(round(j22, 2)) if j22 else ''
if f22 is None or f22 == 0:
    J22_ad = '持平'
    J22_ll = '平'
elif f22 > 0.0:
    J22_ad = '增加'
    J22_ll = '升'
elif f22 < 0.0:
    J22_ad = '减少'
    J22_ll = '降'
k22 = ws.cell(row=22, column=11).value
K22 = "{:,}".format(round(k22, 2)) if k22 else ''

f24 = ws.cell(row=24, column=6).value
F24 = "{:,}".format(round(f24, 2)) if f24 else ''
j24 = ws.cell(row=24, column=10).value
J24 = "{:,}".format(round(j24, 2)) if j24 else ''
if f24 is None or f24 == 0:
    J24_ad = '持平'
    J24_ll = '平'
elif f24 > 0.0:
    J24_ad = '增加'
    J24_ll = '升'
elif f24 < 0.0:
    J24_ad = '减少'
    J24_ll = '降'
k24 = ws.cell(row=24, column=11).value
K24 = "{:,}".format(round(k24, 2)) if k24 else ''

g21 = ws.cell(row=21, column=7).value
G21 = "{:,}".format(round(g21, 2)) if g21 else ''
n21 = ws.cell(row=21, column=14).value
N21 = "{:,}".format(round(n21, 2)) if n21 else ''
if n21 is None or n21 == 0:
    N21_ad = '持平'
    N21_ll = '平'
elif n21 > 0.0:
    N21_ad = '增加'
    N21_ll = '升'
elif n21 < 0.0:
    N21_ad = '减少'
    N21_ll = '降'
o21 = ws.cell(row=21, column=15).value
O21 = "{:,}".format(round(o21, 2)) if o21 else ''

g22 = ws.cell(row=22, column=7).value
G22 = "{:,}".format(round(g22, 2)) if g22 else ''
n22 = ws.cell(row=22, column=14).value
N22 = "{:,}".format(round(n22, 2)) if n22 else ''
if n22 is None or n22 == 0:
    N22_ad = '持平'
    N22_ll = '平'
elif n22 > 0.0:
    N22_ad = '增加'
    N22_ll = '升'
elif n22 < 0.0:
    N22_ad = '减少'
    N22_ll = '降'
o22 = ws.cell(row=22, column=15).value
O22 = "{:,}".format(round(o22, 2)) if o22 else ''

g24 = ws.cell(row=24, column=7).value
G24 = "{:,}".format(round(g24, 2)) if g24 else ''
n24 = ws.cell(row=24, column=14).value
N24 = "{:,}".format(round(n24, 2)) if n24 else ''
if n24 is None or n24 == 0:
    N24_ad = '持平'
    N24_ll = '平'
elif n24 > 0.0:
    N24_ad = '增加'
    N24_ll = '升'
elif n24 < 0.0:
    N24_ad = '减少'
    N24_ll = '降'
o24 = ws.cell(row=24, column=15).value
O24 = "{:,}".format(round(o24, 2)) if o24 else ''
wb.close()

text1 = f'公司当月实现自营业务净收入{F21}{Unit}，环比{J21_ad}{J21}{Unit}，{J21_ll}幅{K21}%；\
信用业务净收入{F22}{Unit}，环比{J22_ad}{J22}{Unit}，{J22_ll}幅{K22}%；\
利润总额{F24}{Unit}，环比{J24_ad}{J24}{Unit}，环比{J24_ll}幅{K24}%。'
text2 = f'1-{date_本年当月月份}月累计实现自营业务净收入{G21}{Unit}，同比{N21_ad}{N21}{Unit}，{N21_ll}幅{O21} %；\
信用业务净收入{G22}{Unit}，同比{N22_ad}{N22}{Unit}，{N22_ll}幅{O22}%；\
利润总额{G24}{Unit}，同比{N24_ad}{N24}{Unit}，同比{N24_ll}幅{O24}%。'

print(text1 + text2)

# -------------------------------------------------word-------------------------------------------------
from docx import Document

doc = Document()
doc.add_paragraph(text1)  # 写入若干段落
doc.add_paragraph(text2)
doc.save("生成变动说明_生成的.docx")  # 保存才能看到结果

DocName = "生成变动说明_生成的.docx"
from win32com.client import Dispatch, constants

word = Dispatch('Word.Application')
word.Visible = True
DocName = os.path.abspath(DocName)  # win32不认识相对路径，故需转换为绝对路径。
doc = word.Documents.Open(DocName)

sleep(5)  # 暂停30秒再关闭word

doc.SaveAs(DocName, 12)
# 关闭word文档
doc.Close()
word.Quit()

# ---------------------------------------辅助余额表-总部-部门--------------------------------------------
wb = load_workbook(xlfile, data_only=True)
ws = wb['辅助余额表-总部-部门']

D05 = ws.cell(row=5, column=4).value
print(D05)
